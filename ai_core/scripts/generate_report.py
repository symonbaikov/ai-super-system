import json, sys, datetime
from docx import Document
from docx.shared import Pt

def make_docx(data, out_path):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    doc.add_heading(f"Отчёт по тикеру — {data['token']}", level=0)
    # Badges line
    p = doc.add_paragraph(f"🔥 {data['hype']}   🐳 {data['whales']}   🛡 {data['risk']}")
    # OG/Clones
    doc.add_heading("OG / Клоны", level=1)
    doc.add_paragraph(f"OG: {data['og']}, клонов: {data['clones']}")
    doc.add_paragraph(f"Оригинальные выражения/картинки: {len(data.get('originals', []))}")
    for o in data.get('originals', []):
        doc.add_paragraph(f"- {o.get('type')}: {o.get('hash')} → {o.get('src')}", style=None)
    # Social
    doc.add_heading("Соц‑эхо", level=1)
    doc.add_paragraph(f"Твиты: {data['tweets']}, Ретвиты: {data['retweets']}, Z-score: {data['zscore']}")
    doc.add_paragraph(f"Первый источник: {data['first_src']}")
    # Whales
    doc.add_heading("Киты/Смарт‑деньги", level=1)
    doc.add_paragraph(f"Китов: {data['whale_count']}, Средний win rate: {data['avg_win_rate']}%, Средний PnL: {data['avg_pnl']}%")
    doc.add_paragraph(f"Держатели 1–7 дней: {data['holder_1_7']}")
    # Risk
    doc.add_heading("Риск‑оценка", level=1)
    doc.add_paragraph(f"Fail‑closed: {data['fail_closed']}, Flags: {', '.join(data.get('flags', []))}")
    # Plan
    doc.add_heading("Торговый план", level=1)
    doc.add_paragraph(f"TP: ×{data['tp']} | SL: {data['sl']}% | TSL: {data['tsl']}%")
    # Footer
    doc.add_paragraph(f"Сгенерировано: {datetime.datetime.utcnow().isoformat()}Z")
    doc.save(out_path)

if __name__ == '__main__':
    inp = sys.argv[1]
    out_docx = sys.argv[2]
    out_html = sys.argv[3]
    html_template = sys.argv[4]
    data = json.load(open(inp,'r',encoding='utf-8'))
    # DOCX
    make_docx(data, out_docx)
    # HTML
    html = open(html_template,'r',encoding='utf-8').read()
    def esc(s): 
        return s.replace('{','{{').replace('}','}}')
    originals_list = "".join([f"<li>{o.get('type')}: {o.get('hash')} → <a href='{o.get('src')}' target='_blank'>{o.get('src')}</a> <small>{o.get('ts')}</small></li>" for o in data.get('originals', [])])
    html_filled = html.replace("{{token}}", data["token"]) \
        .replace("{{hype}}", str(data["hype"])) \
        .replace("{{whales}}", str(data["whales"])) \
        .replace("{{risk}}", data["risk"]) \
        .replace("{{ts}}", datetime.datetime.utcnow().isoformat()+"Z") \
        .replace("{{og}}", str(data["og"])) \
        .replace("{{clones}}", str(data["clones"])) \
        .replace("{{originals_count}}", str(len(data.get("originals", [])))) \
        .replace("{{originals_list}}", originals_list) \
        .replace("{{tweets}}", str(data["tweets"])) \
        .replace("{{retweets}}", str(data["retweets"])) \
        .replace("{{zscore}}", str(data["zscore"])) \
        .replace("{{first_src}}", data.get("first_src","")) \
        .replace("{{whale_count}}", str(data["whale_count"])) \
        .replace("{{avg_win_rate}}", str(data["avg_win_rate"])) \
        .replace("{{avg_pnl}}", str(data["avg_pnl"])) \
        .replace("{{holder_1_7}}", str(data["holder_1_7"])) \
        .replace("{{fail_closed}}", str(data["fail_closed"])) \
        .replace("{{flags}}", ", ".join(data.get("flags", []))) \
        .replace("{{tp}}", str(data["tp"])) \
        .replace("{{sl}}", str(data["sl"])) \
        .replace("{{tsl}}", str(data["tsl"]))
    open(out_html,'w',encoding='utf-8').write(html_filled)
